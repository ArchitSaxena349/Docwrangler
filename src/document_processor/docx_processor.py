from docx import Document
from typing import Dict, Any
from .base import BaseDocumentProcessor

class DocxProcessor(BaseDocumentProcessor):
    """DOCX document processor"""
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
            
        except Exception as e:
            raise Exception(f"Error processing DOCX: {str(e)}")
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata from DOCX"""
        metadata = {"file_type": "docx", "file_path": file_path}
        
        try:
            doc = Document(file_path)
            core_props = doc.core_properties
            
            metadata.update({
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            })
            
        except Exception as e:
            metadata["extraction_error"] = str(e)
        
        return metadata